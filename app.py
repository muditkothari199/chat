import tkinter as tk
from tkinter import filedialog
import streamlit as st
import pandas as pd
import tempfile
import io
import logging
from docx import Document
from docx.shared import RGBColor
import PyPDF2
from pdf2image import convert_from_path
import pytesseract
import cohere
from nltk.corpus import stopwords

# Setup basic logging
logging.basicConfig(level=logging.DEBUG)

# Set the paths for Tesseract and Poppler
pytesseract.pytesseract.tesseract_cmd = r"Tesseract-OCR\tesseract.exe"
poppler_path = r"poppler-24.07.0\Library\bin"

# Initialize Cohere API client
co = cohere.Client('0VBE3VyqTeKVmbxEYE2t9sOEKroR4ZG11SrD6v6M')  # Replace with your actual Cohere API key

# Define the function to query the Cohere API using `co.chat`
def query_cohere_api(text, query):
    combined_message = f"Document Text: {text}\nUser Query: {query}"
    try:
        response = co.chat(
            message=combined_message,
            model="command-r-plus",
            temperature=0
        )
        return response.text if hasattr(response, 'text') else "No content available"
    except Exception as e:
        logging.error(f"API Request Error: {e}")
        return "Error"

# Define a function to remove stopwords from text
def remove_stopwords(text):
    stop_words = set(stopwords.words('english'))
    filtered_words = [word for word in text.split() if word.lower() not in stop_words]
    return ' '.join(filtered_words)

# Define a function to extract text from both searchable and scanned PDFs using OCR
def extract_text_from_pdf(pdf_path, max_pages=15):
    text = ""
    try:
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            num_pages = min(max_pages, len(reader.pages))  # Limit to max_pages or total pages, whichever is smaller
            for page_num in range(num_pages):
                page = reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text += page_text
                else:
                    images = convert_from_path(pdf_path, poppler_path=poppler_path, first_page=page_num+1, last_page=page_num+1)
                    for image in images:
                        ocr_text = pytesseract.image_to_string(image)
                        text += ocr_text
        text = remove_stopwords(text)
    except Exception as e:
        logging.error(f"Failed to extract text from PDF: {e}")
    return text

# Define a function to handle the special case for "HO-3"
def handle_ho3_policy_type(result):
    if "HO-3" in result.upper() or "HO3" in result.upper():
        return "Homeowners"
    return result

# Define a function to process PDFs and return results in a dictionary
def process_pdfs(pdf_paths):
    queries = [
        'Provide the Full named insured in one word.',
        'only the policy period or policy date concisely.',
        'Give the insurance company name only.',
        'only the name of coverages and only their limits of liability to the point answer.',
        'Give me the name of policy type.'
    ]

    short_names = {
        'Provide the Full named insured in one word.': 'Named Insured',
        'only the policy period or policy date concisely.': 'Policy Date',
        'Give the insurance company name only.': 'Insurance Company',
        'only the name of coverages and only their limits of liability to the point answer.': 'Coverage & Limits',
        'Give me the name of policy type.': 'Policy Type'
    }

    key_list = [
        'Automobile Policy', 'Home', 'Auto', 'Homeowners Insurance', 'Renters Insurance', 'Flood Insurance',
        'Condo Insurance', 'Mobile Home Insurance', 'Travel Insurance', 'Disability Insurance',
        'Pet Insurance', 'Earthquake Insurance', 'Umbrella Insurance', "Workers' Compensation Insurance",
        'Cyber Liability Insurance', 'Business Insurance', 'General Liability Insurance',
        'Property Insurance', 'Business Interruption Insurance', 'Professional Liability Insurance',
        'Commercial Multiple Peril Policies', 'Term Life Insurance', 'Whole Life Insurance',
        'Product Liability Insurance', 'Builders Risk Insurance', 'Commercial Crime Insurance',
        'Environmental Liability Insurance', 'Marine Insurance', 'Universal Life Insurance',
        'Variable Life Insurance', 'Indexed Universal Life Insurance', 'Final Expense Insurance',
        'Survivorship Life Insurance (Second-to-Die Insurance)', 'Guaranteed Issue Life Insurance',
        'No-Exam Life Insurance', 'Health Maintenance Organization (HMO)',
        'Preferred Provider Organization (PPO)', 'Exclusive Provider Organization (EPO)',
        'Point of Service (POS)', 'High Deductible Health Plan (HDHP)', 'Catastrophic Health Insurance',
        'Medicare Part A (Hospital Insurance)', 'Medicare Part B (Medical Insurance)',
        'Medicare Part C (Medicare Advantage)', 'Medicare Part D (Prescription Drug Coverage)',
        'Childrens Health Insurance Program (CHIP)', 'Individual Health Insurance',
        'Group Health Insurance', 'Medicaid', 'Short-Term Health Insurance',
        'Critical Illness Insurance', 'Accident Insurance', 'Dental Insurance', 'Vision Insurance',
        'Long-Term Care Insurance', 'Health Savings Accounts (HSAs) and Flexible Spending Accounts (FSAs)',
        'Cyber Insurance', 'Environmental Liability Insurance',
        'Professional Liability Insurance (Errors and Omissions Insurance)', 'Medical Malpractice Insurance',
        'Legal Malpractice Insurance', 'Financial Services Malpractice Insurance',
        'Directors and Officers (D&O) Insurance', 'Kidnap and Ransom Insurance', 'Terrorism Insurance',
        'Event Cancellation Insurance', 'Fine Arts Insurance', 'Excess and Surplus Lines Insurance', "Auto",
        "Cyber Liability", "Directors & Officers", "Employment Practices Liability", "Errors & Omissions", "Excess Liability",
        "General Liability", "Homeowners", "Individual Health", "Inland Marine", "Professional Liability", "Property", "Umbrella",
        "Flood", "Dwelling Fire", "Mobile Home", "Businessowners Policy", "Boat", "Property", "Earthquake", "Garage", "Fire", "Bond", "Workers Compensation"
    ]

    results_dict = {}

    for index, file_path in enumerate(pdf_paths):
        doc_key = f'doc{index+1}'
        results_dict[doc_key] = {short_names.get(query, query): '' for query in queries}

        if file_path:
            logging.info(f"Processing file: {file_path}")
            all_text = extract_text_from_pdf(file_path)
            logging.debug(f"Extracted Text from {file_path}: {all_text[:500]}")  # Print the first 500 characters of extracted text for debugging
            
            for query in queries:
                result = query_cohere_api(all_text, query)
                if query == queries[-1]:  # Only check key_list and HO-3 for the last query (Policy Type)
                    result = handle_ho3_policy_type(result)
                    if any(policy_type.lower() in result.lower() for policy_type in key_list):
                        matching_policy = next(policy_type for policy_type in key_list if policy_type.lower() in result.lower())
                        results_dict[doc_key][short_names[query]] = matching_policy
                    else:
                        results_dict[doc_key][short_names[query]] = result
                else:
                    results_dict[doc_key][short_names[query]] = result

    save_to_word_and_prompt(results_dict, short_names)
    return results_dict

# Modify this function to save results of multiple PDFs into one Word document
def save_to_word(results_dict, short_names):
    output = io.BytesIO()
    document = Document()
    document.add_heading('PDF Processing Results', level=1)

    table = document.add_table(rows=1, cols=len(results_dict) + 1)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Query'

    # Add document columns
    for i, doc_key in enumerate(results_dict.keys()):
        hdr_cells[i + 1].text = f'{doc_key}'

    

    # Add rows for each query
    for query in short_names.values():
        row_cells = table.add_row().cells
        row_cells[0].text = query
        for i, doc_key in enumerate(results_dict.keys()):
            row_cells[i + 1].text = results_dict[doc_key].get(query, '')



    document.save(output)
    return output.getvalue()  # Return the Word file as a byte stream

def save_to_word_and_prompt(results_dict, short_names):
    word_file_bytes = save_to_word(results_dict, short_names)
    
    # Create a new Tkinter root window
    root = tk.Tk()
    root.withdraw()  # Hide the main window

    try:
        save_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
            title="Save Combined Document"
        )
        if save_path:
            with open(save_path, "wb") as f:
                f.write(word_file_bytes)
    finally:
        root.destroy()  # Destroy the Tkinter root window to ensure it closes properly

# Streamlit interface
st.title("PDF Processor")

# Use the sidebar for file upload
uploaded_files = st.sidebar.file_uploader("Upload PDF files", accept_multiple_files=True, type="pdf")

if uploaded_files:
    pdf_paths = []
    for uploaded_file in uploaded_files:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
            tmp_file.write(uploaded_file.read())
            pdf_paths.append(tmp_file.name)
    
    if st.button("Process PDFs"):
       # st.write("Processing...")
        results_dict = process_pdfs(pdf_paths)
        st.write("Processing Complete!")

        # Display results in a table
        df = pd.DataFrame(results_dict)
        st.table(df)


if __name__ == "__main__":
    st.write("Streamlit app running")
